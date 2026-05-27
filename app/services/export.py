from pathlib import Path
from typing import Iterable, List, Optional

from app.config import EXPORTS_DIR
from app.models import DocumentBlock, TranslationJob

# 한글(CJK)을 지원하는 시스템 폰트 후보 목록 (macOS → Linux → Windows 순)
_CJK_FONT_CANDIDATES = [
    "/System/Library/Fonts/AppleSDGothicNeo.ttc",          # macOS 기본 한글 폰트
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",  # macOS 구버전
    "/Library/Fonts/NanumGothic.ttf",                      # 나눔고딕 (설치 시)
    "/Library/Fonts/NanumBarunGothic.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansCJKkr-Regular.otf",  # Linux Noto
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",     # Linux 나눔
    "/usr/share/fonts/truetype/unfonts-core/UnDotum.ttf",  # Linux 은폰트
    "C:/Windows/Fonts/malgun.ttf",                         # Windows 맑은 고딕
    "C:/Windows/Fonts/gulim.ttc",
]


def _find_cjk_font() -> Optional[str]:
    """시스템에서 한글을 지원하는 폰트 파일 경로를 반환. 없으면 None."""
    for p in _CJK_FONT_CANDIDATES:
        if Path(p).exists():
            return p
    return None


def export_job(job: TranslationJob, blocks: List[DocumentBlock]) -> Path:
    if job.outputFormat in {"txt", "md"}:
        return export_text(job, blocks, job.outputFormat)
    if job.outputFormat == "pdf":
        # bbox가 있는 블록(pdf_text_span / pdf_ocr_span)이 하나라도 있으면
        # 원위치 교체(layout-preserving) 방식을 우선 사용.
        # 모든 블록이 pdf_vision_span(bbox 없음)일 때만 하단 오버레이 방식으로 폴백.
        has_inplace = any(b.type in ("pdf_text_span", "pdf_ocr_span") for b in blocks)
        if has_inplace:
            return export_pdf_layout_preserving(job, blocks)
        return export_pdf_vision_ocr(job, blocks)
    if job.outputFormat in {"doc", "docx"}:
        return export_docx_preserving(job, blocks)
    if job.outputFormat in {"xls", "xlsx"}:
        return export_xlsx_preserving(job, blocks)
    raise ValueError(f"{job.outputFormat.upper()} 내보내기는 지원하지 않습니다")


def export_text(job: TranslationJob, blocks: Iterable[DocumentBlock], extension: str) -> Path:
    out_path = EXPORTS_DIR / f"{job.id}.{extension}"
    text = "\n\n".join((block.translatedText or "") for block in blocks)
    out_path.write_text(text, encoding="utf-8")
    return out_path


def export_pdf_layout_preserving(job: TranslationJob, blocks: List[DocumentBlock]) -> Path:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "Layout-preserving PDF export requires PyMuPDF. Install dependencies with: python3 -m pip install PyMuPDF"
        ) from exc

    if not job.originalPath:
        raise RuntimeError("Original PDF path is missing")

    out_path = EXPORTS_DIR / f"{job.id}.pdf"
    doc = fitz.open(job.originalPath)

    # pdf_text_span: 벡터 텍스트 교체 / pdf_ocr_span: 이미지 위 오버레이
    pdf_blocks = [
        block for block in blocks
        if block.type in ("pdf_text_span", "pdf_ocr_span") and block.bbox
    ]

    # ── ① 벡터 텍스트(pdf_text_span)만 Redact으로 제거 ──────────────────
    # (pdf_ocr_span은 이미지 위에 있으므로 redact 불필요 — draw_rect로 처리)
    for block in pdf_blocks:
        if block.type != "pdf_text_span":
            continue
        page = doc[block.pageNumber - 1 if block.pageNumber else 0]
        rect = fitz.Rect(*block.bbox)
        page.add_redact_annot(rect, fill=(1, 1, 1))

    for page in doc:
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    # ── ② 한글 폰트 준비 ─────────────────────────────────────────────────
    cjk_font_path = _find_cjk_font()
    font_registered: set = set()   # 페이지별 폰트 등록 여부 추적

    # ── ③ 번역 텍스트 삽입 ───────────────────────────────────────────────
    for block in pdf_blocks:
        page_idx = block.pageNumber - 1 if block.pageNumber else 0
        page = doc[page_idx]

        # 페이지당 한 번만 한글 폰트 등록
        if cjk_font_path and page_idx not in font_registered:
            try:
                page.insert_font(fontname="KR", fontfile=cjk_font_path)
                font_registered.add(page_idx)
            except Exception:
                pass  # 폰트 등록 실패 시 기본 폰트로 fallback

        fontname = "KR" if page_idx in font_registered else "helv"
        rect = fitz.Rect(*block.bbox)
        text = block.translatedText or ""
        font_size = max(5, min(block.fontSize or 9, 14))

        # OCR 블록: 이미지 위에 흰 사각형을 덮어 원본 텍스트 이미지를 가림
        if block.type == "pdf_ocr_span":
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))

        inserted = page.insert_textbox(
            rect, text, fontsize=font_size, fontname=fontname,
            color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT,
        )
        # 공간이 부족하면 폰트 크기를 줄여서 재시도
        while inserted < 0 and font_size > 5:
            font_size -= 0.5
            inserted = page.insert_textbox(
                rect, text, fontsize=font_size, fontname=fontname,
                color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT,
            )

    doc.save(out_path)
    doc.close()
    return out_path


def export_pdf_vision_ocr(job: TranslationJob, blocks: List[DocumentBlock]) -> Path:
    """Claude Vision OCR 블록(bbox 없음)을 번역된 텍스트 PDF로 내보내기.

    레이아웃:
      ┌──────────────────────┐
      │  원본 페이지 (래스터) │  ← 1x zoom (72dpi), 폰트 참조 없음
      ├──────────────────────┤
      │  번역 텍스트 (흰 배경)│  ← 전체 폭, 텍스트 양에 따라 높이 자동 조정
      └──────────────────────┘

    파일 크기 최적화:
      - zoom=1.0 (원본 해상도 유지, 불필요한 확대 없음)
      - garbage=4 + deflate=True 로 저장 → 폰트 중복 제거 + 압축
    """
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyMuPDF가 필요합니다") from exc

    import logging as _logging
    from collections import defaultdict

    _log = _logging.getLogger(__name__)
    _FONT_SIZE   = 10
    _PADDING     = 14
    _LINE_HEIGHT = _FONT_SIZE * 1.55
    _MIN_TRANS_H = 100

    out_path = EXPORTS_DIR / f"{job.id}.pdf"
    cjk_font_path = _find_cjk_font()

    # 폰트 데이터를 한 번만 읽어 모든 페이지에서 재사용 (중복 임베딩 방지)
    cjk_font_data: Optional[bytes] = None
    if cjk_font_path:
        try:
            cjk_font_data = Path(cjk_font_path).read_bytes()
        except Exception:
            cjk_font_data = None

    orig_doc = None
    if job.originalPath and Path(job.originalPath).exists():
        orig_doc = fitz.open(job.originalPath)

    out_doc = fitz.open()

    # ── 페이지별 블록 그룹화 ────────────────────────────────────────────
    page_blocks: dict = defaultdict(list)
    for block in blocks:
        if block.type == "pdf_vision_span":
            page_blocks[block.pageNumber or 1].append(block)

    page_nums = sorted(page_blocks.keys())
    if not page_nums:
        out_doc.new_page()
        out_doc.save(out_path, garbage=4, deflate=True)
        out_doc.close()
        if orig_doc:
            orig_doc.close()
        return out_path

    for page_num in page_nums:
        blks = page_blocks[page_num]

        # ── 원본 페이지 크기 ──────────────────────────────────────────────
        if orig_doc and page_num - 1 < len(orig_doc):
            orig_page = orig_doc[page_num - 1]
            w = float(orig_page.rect.width)
            h = float(orig_page.rect.height)
        else:
            orig_page = None
            w, h = 595.0, 842.0

        # ── 번역 텍스트 수집 ─────────────────────────────────────────────
        combined_text = "\n\n".join(
            b.translatedText for b in blks if b.translatedText and b.translatedText.strip()
        )

        # ── 번역 영역 높이 추정 ───────────────────────────────────────────
        usable_w     = w - 2 * _PADDING
        avg_char_w   = _FONT_SIZE * 0.72   # 한글 글자폭 근사
        chars_per_ln = max(1, int(usable_w / avg_char_w))
        para_list    = [p for p in combined_text.split("\n\n") if p.strip()] if combined_text else []
        total_lines  = sum(max(1, (len(p) // chars_per_ln) + 1) for p in para_list)
        trans_h      = max(_MIN_TRANS_H,
                           total_lines * _LINE_HEIGHT + len(para_list) * 6 + 2 * _PADDING)

        new_h    = h + trans_h
        new_page = out_doc.new_page(width=w, height=new_h)

        # ── 원본 페이지 → 래스터 이미지 삽입 (zoom=1.0, 폰트 참조 없음) ──
        if orig_page is not None:
            try:
                pix = orig_page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)
                new_page.insert_image(fitz.Rect(0, 0, w, h), pixmap=pix)
            except Exception as exc:
                _log.warning("[export] p%d 원본 이미지 삽입 실패: %s", page_num, exc)

        # ── 번역 영역 배경 + 구분선 ──────────────────────────────────────
        new_page.draw_rect(
            fitz.Rect(0, h, w, new_h),
            color=(0.8, 0.8, 0.8), fill=(1, 1, 1),
        )
        new_page.draw_line(
            fitz.Point(0, h), fitz.Point(w, h),
            color=(0.4, 0.4, 0.4), width=0.8,
        )

        if not combined_text:
            _log.warning("[export] p%d 번역 텍스트 없음 (translatedText=None)", page_num)
            continue

        # ── 폰트 등록 (페이지마다 등록, garbage=4 저장 시 자동 중복 제거) ─
        fontname = "helv"
        if cjk_font_data:
            try:
                new_page.insert_font(fontname="KR", fontbuffer=cjk_font_data)
                fontname = "KR"
            except Exception as exc:
                _log.warning("[export] p%d KR 폰트 등록 실패: %s", page_num, exc)

        # ── 텍스트 삽입 ──────────────────────────────────────────────────
        inner_rect = fitz.Rect(_PADDING, h + _PADDING, w - _PADDING, new_h - _PADDING)
        try:
            overflow = new_page.insert_textbox(
                inner_rect, combined_text,
                fontsize=_FONT_SIZE, fontname=fontname,
                color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT,
            )
            # 넘치면 더 작은 폰트로 재시도
            if overflow < 0 and _FONT_SIZE > 7:
                new_page.insert_textbox(
                    inner_rect, combined_text,
                    fontsize=max(7, _FONT_SIZE - 2), fontname=fontname,
                    color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT,
                )
        except Exception as exc:
            _log.warning("[export] p%d 텍스트 삽입 실패(%s): %s", page_num, fontname, exc)
            # CJK 폰트 실패 시 helv 로 재시도 (글자 깨질 수 있지만 페이지 유지)
            if fontname != "helv":
                try:
                    new_page.insert_textbox(
                        inner_rect, combined_text,
                        fontsize=_FONT_SIZE, fontname="helv",
                        color=(0, 0, 0), align=fitz.TEXT_ALIGN_LEFT,
                    )
                except Exception:
                    pass

    if orig_doc:
        orig_doc.close()

    # garbage=4: 중복 폰트·이미지 제거 / deflate: 스트림 압축
    out_doc.save(out_path, garbage=4, deflate=True)
    out_doc.close()
    return out_path


def export_docx_preserving(job: TranslationJob, blocks: List[DocumentBlock]) -> Path:
    """원본 docx의 스타일/서식을 유지하면서 번역 텍스트만 교체.
    extraction.py 의 extract_docx_blocks()와 block.id 구조를 공유합니다.
    원본이 .doc 바이너리인 경우 서식 없이 새 DOCX를 생성합니다."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "DOCX 내보내기에는 python-docx가 필요합니다: pip install python-docx"
        ) from exc

    if not job.originalPath:
        raise RuntimeError("원본 파일 경로가 없습니다")

    original_ext = Path(job.originalPath).suffix.lower()

    # ── .doc 바이너리: python-docx로 열 수 없으므로 새 문서로 생성 ──
    if original_ext == ".doc":
        return _export_docx_from_blocks(blocks, EXPORTS_DIR / f"{job.id}.docx")

    # ── .docx: 원본 열고 텍스트만 교체 (서식 유지) ───────────────────
    doc = Document(job.originalPath)
    translation_map = {b.id: (b.translatedText or "") for b in blocks}

    # 일반 단락 교체
    for para_idx, para in enumerate(doc.paragraphs):
        translated = translation_map.get(f"para-{para_idx}")
        if translated is None:
            continue
        # runs 전체를 유지하되 첫 번째 run에 번역문을 넣고 나머지를 비움
        # → bold/italic/font/color 등 인라인 서식이 첫 run 기준으로 유지됨
        if para.runs:
            para.runs[0].text = translated
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.clear()
            para.add_run(translated)

    # 표 셀 교체
    seen_cells: set = set()
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                translated = translation_map.get(f"tbl-{tbl_idx}-r{row_idx}-c{col_idx}")
                if translated is None:
                    continue
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = translated
                    for run in cell.paragraphs[0].runs[1:]:
                        run.text = ""
                elif cell.paragraphs:
                    cell.paragraphs[0].clear()
                    cell.paragraphs[0].add_run(translated)

    out_path = EXPORTS_DIR / f"{job.id}.docx"
    doc.save(out_path)
    return out_path


def _export_docx_from_blocks(blocks: List[DocumentBlock], out_path: Path) -> Path:
    """번역 블록만으로 새 DOCX 문서를 생성 (서식 없음).
    .doc 바이너리 원본처럼 원본 서식을 읽을 수 없는 경우에 사용."""
    from docx import Document  # type: ignore
    doc = Document()
    for block in blocks:
        text = block.translatedText or ""
        if not text.strip():
            continue
        doc.add_paragraph(text)
    doc.save(out_path)
    return out_path


def export_xlsx_preserving(job: TranslationJob, blocks: List[DocumentBlock]) -> Path:
    """xlsx 서식을 유지하면서 번역 텍스트만 교체.
    원본이 .xls 바이너리인 경우 서식 없이 새 XLSX를 생성합니다."""
    try:
        import openpyxl  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "XLSX 내보내기에는 openpyxl이 필요합니다: pip install openpyxl"
        ) from exc

    if not job.originalPath:
        raise RuntimeError("원본 파일 경로가 없습니다")

    out_path = EXPORTS_DIR / f"{job.id}.xlsx"
    original_ext = Path(job.originalPath).suffix.lower()

    # ── .xls 바이너리: openpyxl로 열 수 없으므로 새 문서로 생성 ───────
    if original_ext == ".xls":
        return _export_xlsx_from_blocks(blocks, out_path)

    # ── .xlsx: 원본 열고 셀 값만 교체 (서식 유지) ────────────────────
    # keep_vba=True 로 매크로 유지, data_only=False 로 수식 셀도 보존
    wb = openpyxl.load_workbook(job.originalPath, keep_vba=True)
    translation_map = {b.id: (b.translatedText or "") for b in blocks}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                block_id = f"sheet-{sheet_name}-{cell.coordinate}"
                translated = translation_map.get(block_id)
                if translated is None:
                    continue
                # 셀 값만 교체 — number_format, font, fill, border 등은 그대로
                cell.value = translated

    wb.save(out_path)
    return out_path


def _export_xlsx_from_blocks(blocks: List[DocumentBlock], out_path: Path) -> Path:
    """번역 블록만으로 새 XLSX 문서를 생성 (서식 없음).
    .xls 바이너리처럼 원본 서식을 읽을 수 없는 경우에 사용."""
    import openpyxl  # type: ignore
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    sheets: dict = {}

    for block in blocks:
        if block.type != "xlsx_cell":
            continue
        sheet_name = block.sheetName or "Sheet1"
        if sheet_name not in sheets:
            if not sheets:
                ws = wb.active
                ws.title = sheet_name
            else:
                ws = wb.create_sheet(sheet_name)
            sheets[sheet_name] = ws
        else:
            ws = sheets[sheet_name]

        coord = block.metadata.get("coordinate", "A1")
        ws[coord] = block.translatedText or ""

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)
    return out_path
