import base64
import logging
import os
import re
import subprocess
from io import BytesIO
from pathlib import Path
from typing import List

from app.models import DocumentBlock

logger = logging.getLogger(__name__)

# ── OCR 설정 ────────────────────────────────────────────────────────────
# 페이지당 추출된 문자 수가 이 값 미만이면 스캔 페이지로 판단 → OCR 시도
_OCR_CHAR_THRESHOLD = 20
# OCR 렌더링 배율 (높을수록 정확도 ↑, 메모리 ↑)
_OCR_ZOOM = 2.0
# Tesseract 최소 신뢰도 (0~100). 이하 단어는 무시
_OCR_MIN_CONF = 30
# Tesseract 언어팩: 일본어·영어·중국어(간체)·중국어(번체)·한국어
_OCR_LANG = "jpn+eng+chi_sim+chi_tra+kor"
# Claude Vision OCR 모델 (translation.py와 동일한 haiku 모델 사용)
_VISION_MODEL = "claude-haiku-4-5-20251001"
# Claude document API 1회 호출당 처리할 최대 페이지 수.
# 고밀도 일본어 기준 페이지당 ~2,000 토큰 → 2페이지 = ~4,000 토큰 (max_tokens 8192 이내)
_PDF_BATCH_PAGES = 2


def extract_blocks(path: Path, file_type: str, ocr_engine: str = "none") -> List[DocumentBlock]:
    if file_type in {"txt", "md"}:
        return extract_text_blocks(path, file_type)
    if file_type == "pdf":
        return extract_pdf_blocks(path, ocr_engine=ocr_engine)
    if file_type == "docx":
        return extract_docx_blocks(path)
    if file_type == "doc":
        return extract_doc_blocks(path)
    if file_type == "xlsx":
        return extract_xlsx_blocks(path)
    if file_type == "xls":
        return extract_xls_blocks(path)
    raise ValueError(f"{file_type.upper()} extraction is not supported")


def extract_text_blocks(path: Path, file_type: str) -> List[DocumentBlock]:
    text = path.read_text(encoding="utf-8", errors="replace")
    parts = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if not parts and text.strip():
        parts = [text.strip()]
    blocks = []
    for index, part in enumerate(parts):
        block_type = "paragraph"
        if file_type == "md" and part.startswith("#"):
            block_type = "heading"
        blocks.append(
            DocumentBlock(
                id=f"block-{index + 1}",
                type=block_type,
                sourceText=part,
                order=index,
                metadata={"fileType": file_type},
            )
        )
    return blocks


def extract_pdf_blocks(path: Path, ocr_engine: str = "none") -> List[DocumentBlock]:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "PDF extraction requires PyMuPDF. Install dependencies with: python3 -m pip install PyMuPDF"
        ) from exc

    doc = fitz.open(path)
    page_count = len(doc)
    blocks: List[DocumentBlock] = []
    order = 0

    # ── Claude document API 사전 추출 (claude_vision 모드) ─────────────────
    # Claude Chat과 동일하게 PDF 전체를 document API로 전송 → 고품질 OCR
    doc_page_texts: dict = {}
    if ocr_engine == "claude_vision":
        try:
            doc_page_texts = _fetch_pdf_claude_doc(path, page_count)
            logger.warning("[PDF Doc API] %d/%d 페이지 사전 추출 완료", len(doc_page_texts), page_count)
        except Exception as exc:
            logger.warning(
                "[PDF Doc API] 사전 추출 실패, 페이지별 Vision API로 폴백: %s", exc
            )

    for page_index, page in enumerate(doc):
        raw = page.get_text("dict")

        # ── 일반 텍스트 추출 ────────────────────────────────────────────
        page_text_blocks: List[DocumentBlock] = []
        for raw_block in raw.get("blocks", []):
            if raw_block.get("type") != 0:
                continue
            for line in raw_block.get("lines", []):
                spans = [span for span in line.get("spans", []) if span.get("text", "").strip()]
                if not spans:
                    continue
                text = "".join(span.get("text", "") for span in spans).strip()
                first_span = spans[0]
                x0, y0, x1, y1 = line.get("bbox", first_span.get("bbox"))
                page_text_blocks.append(
                    DocumentBlock(
                        id=f"page-{page_index + 1}-span-{order + len(page_text_blocks) + 1}",
                        type="pdf_text_span",
                        sourceText=text,
                        translatedText=None,
                        pageNumber=page_index + 1,
                        bbox=(float(x0), float(y0), float(x1), float(y1)),
                        fontSize=float(first_span.get("size", 10)),
                        order=order + len(page_text_blocks),
                        metadata={
                            "font": first_span.get("font"),
                            "color": first_span.get("color"),
                        },
                    )
                )

        # ── 스캔 페이지 감지 → OCR 분기 ──────────────────────────────────
        page_char_count = sum(len(b.sourceText) for b in page_text_blocks)
        is_scanned = page_char_count < _OCR_CHAR_THRESHOLD

        if is_scanned and ocr_engine in {"tesseract", "claude_vision"}:
            logger.warning(
                "[OCR] 페이지 %d 스캔 감지 (chars=%d) → %s",
                page_index + 1, page_char_count, ocr_engine,
            )
            try:
                page_num = page_index + 1
                if ocr_engine == "tesseract":
                    ocr_blocks = _extract_page_ocr(page, page_index, order)
                elif page_num in doc_page_texts and doc_page_texts[page_num]:
                    # 사전 추출 텍스트 있음 → Tesseract bbox에 매핑
                    logger.warning("[OCR] 페이지 %d: document API 텍스트 사용", page_num)
                    ocr_blocks = _map_claude_text_to_bboxes(
                        page, page_index, order, doc_page_texts[page_num]
                    )
                elif page_num in doc_page_texts:
                    # 사전 추출에서 blank로 확인된 페이지 → Vision API 불필요
                    logger.warning("[OCR] 페이지 %d: document API 확인 공백 페이지, 건너뜀", page_num)
                    ocr_blocks = []
                else:
                    # 배치 실패로 미처리 → 페이지별 Vision API 폴백
                    logger.warning("[OCR] 페이지 %d: 배치 미처리, Vision API 폴백", page_num)
                    ocr_blocks = _extract_page_claude_vision(page, page_index, order)

                if ocr_blocks:
                    blocks.extend(ocr_blocks)
                    order += len(ocr_blocks)
                    continue
                else:
                    # 결과 없음 = blank 페이지 또는 OCR 실패 → 다음 페이지로
                    logger.warning("[OCR] 페이지 %d: OCR 결과 없음 (blank 또는 실패)", page_index + 1)
                    continue
            except Exception as exc:
                logger.warning("[OCR] 페이지 %d 처리 실패: %s", page_index + 1, exc)

        blocks.extend(page_text_blocks)
        order += len(page_text_blocks)

    doc.close()
    return blocks


def _fetch_pdf_claude_doc(path: Path, page_count: int) -> dict:
    """PDF를 _PDF_BATCH_PAGES 페이지씩 나눠 Claude document API로 추출.

    전체를 한 번에 보내면 max_tokens(8192) 한계로 중간에 잘림.
    배치 분할로 안정적으로 전체 페이지를 추출.

    Returns: {page_num(1-based): text_str}
    """
    import json
    import urllib.request

    import fitz  # type: ignore

    from app.provider_settings import provider_model as _pm, provider_secret

    api_key = provider_secret("anthropic", "api_key", "ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("Anthropic API 키 미설정 (config/providers.local.json)")
    vision_model = _pm("anthropic", "model", "ANTHROPIC_TRANSLATION_MODEL", _VISION_MODEL)

    all_page_texts: dict = {}
    src_doc = fitz.open(path)

    try:
        for batch_start in range(0, page_count, _PDF_BATCH_PAGES):
            batch_end = min(batch_start + _PDF_BATCH_PAGES, page_count)
            batch_size = batch_end - batch_start

            # 배치 페이지들을 별도 PDF 바이트로 추출
            sub_doc = fitz.open()
            sub_doc.insert_pdf(src_doc, from_page=batch_start, to_page=batch_end - 1)
            pdf_bytes = sub_doc.tobytes()
            sub_doc.close()

            pdf_b64 = base64.standard_b64encode(pdf_bytes).decode()

            prompt = (
                f"This PDF has {batch_size} page(s). "
                "Extract all text from every page. "
                "For each page output exactly '=== PAGE N ===' "
                "(where N is the page number starting from 1) "
                "on its own line, then the full text of that page. "
                "Rules:\n"
                "- Output ONLY the raw text and page markers, no explanations\n"
                "- Preserve the original language exactly (Japanese, Chinese, Korean, etc.)\n"
                "- Preserve all numbers, punctuation, and special characters\n"
                "- For vertical Japanese/Chinese text: read each column left-to-right, top-to-bottom\n"
                "- Separate distinct text blocks within a page with a blank line\n"
                "- Do NOT translate, summarize, or add markdown formatting\n"
                "- If a page is blank, write '=== PAGE N ===' then '(blank)'"
            )

            payload = json.dumps({
                "model": vision_model,
                "max_tokens": 8192,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_b64,
                            },
                        },
                        {"type": "text", "text": prompt},
                    ],
                }],
            }).encode("utf-8")

            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=payload,
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "anthropic-beta": "pdfs-2024-09-25",
                },
            )

            try:
                with urllib.request.urlopen(req, timeout=120) as resp:
                    body = json.loads(resp.read().decode("utf-8"))
            except Exception as exc:
                logger.warning(
                    "[PDF Doc API] 배치 p%d~p%d 실패: %s",
                    batch_start + 1, batch_end, exc,
                )
                continue

            raw_text = body["content"][0]["text"].strip()
            stop_reason = body.get("stop_reason", "")
            if stop_reason == "max_tokens":
                logger.warning(
                    "[PDF Doc API] 배치 p%d~p%d: max_tokens 도달 — 배치 크기를 줄이세요",
                    batch_start + 1, batch_end,
                )

            logger.warning(
                "[PDF Doc API] 배치 p%d~p%d 추출: %d chars (stop=%s)",
                batch_start + 1, batch_end, len(raw_text), stop_reason,
            )

            # 로컬 페이지 번호(1-based) → 전역 페이지 번호로 변환 후 저장
            # blank 페이지도 ""로 저장 → "배치 실패(키 없음)"와 구분
            marker_re = re.compile(r"===\s*PAGE\s+(\d+)\s*===", re.IGNORECASE)
            parts = marker_re.split(raw_text)
            i = 1
            while i + 1 < len(parts):
                try:
                    local_page = int(parts[i])
                except ValueError:
                    i += 2
                    continue
                text = parts[i + 1].strip()
                global_page = batch_start + local_page
                # blank → "" 저장, 내용 있으면 그대로 저장
                all_page_texts[global_page] = "" if text.lower() in {"(blank)", ""} else text
                i += 2

            # API 응답에 마커가 전혀 없는 경우 (1페이지 배치 등) 안전 처리
            if not any(batch_start < p <= batch_end for p in all_page_texts):
                # 파싱 실패 시 전체 raw_text를 첫 번째 페이지 텍스트로 사용
                if raw_text and raw_text.lower() not in {"(blank)", ""}:
                    all_page_texts[batch_start + 1] = raw_text
                    logger.warning(
                        "[PDF Doc API] 배치 p%d~p%d: PAGE 마커 없음, 전체 텍스트를 p%d에 할당",
                        batch_start + 1, batch_end, batch_start + 1,
                    )

    finally:
        src_doc.close()

    logger.warning(
        "[PDF Doc API] 전체 완료: %d/%d 페이지 추출됨",
        len(all_page_texts), page_count,
    )
    return all_page_texts


def _map_claude_text_to_bboxes(
    page, page_index: int, order_start: int, claude_text: str
) -> List[DocumentBlock]:
    """Claude document API로 추출한 텍스트를 Tesseract bbox에 매핑.

    1. Tesseract로 단락 bbox 탐지 (위치 정보)
    2. claude_text를 단락 단위로 분할
    3. bbox ↔ 텍스트 비례 매핑 → pdf_ocr_span 블록 반환
    Tesseract 미설치 시: pdf_vision_span으로 폴백 (bbox 없음)
    """
    para_bboxes = _get_paragraph_bboxes(page)

    cv_paras = [p.strip() for p in re.split(r"\n{2,}", claude_text) if p.strip()]
    if not cv_paras:
        cv_paras = [ln.strip() for ln in claude_text.splitlines() if ln.strip()]
    if not cv_paras:
        cv_paras = [claude_text] if claude_text else []

    # Tesseract bbox 없으면 pdf_vision_span으로 폴백
    if not para_bboxes:
        result: List[DocumentBlock] = []
        for idx, para in enumerate(cv_paras):
            result.append(DocumentBlock(
                id=f"page-{page_index + 1}-vision-{order_start + idx + 1}",
                type="pdf_vision_span",
                sourceText=para,
                translatedText=None,
                pageNumber=page_index + 1,
                bbox=None,
                order=order_start + idx,
                metadata={"ocr_engine": "claude_doc", "ocr_confidence": 100},
            ))
        return result

    # Tesseract 단락 수(n)에 맞게 Claude 단락(m)을 균등 분배
    n, m = len(para_bboxes), len(cv_paras)
    out: List[DocumentBlock] = []
    for i, (x0, y0, x1, y1, _) in enumerate(para_bboxes):
        start = round(i * m / n)
        end   = round((i + 1) * m / n)
        text  = " ".join(cv_paras[start:end]).strip()
        if not text:
            continue
        font_size = max(6.0, round((y1 - y0) * 0.75, 1))
        out.append(DocumentBlock(
            id=f"page-{page_index + 1}-doc-{order_start + i + 1}",
            type="pdf_ocr_span",        # bbox 있음 → 원위치 교체
            sourceText=text,
            translatedText=None,
            pageNumber=page_index + 1,
            bbox=(x0, y0, x1, y1),
            fontSize=font_size,
            order=order_start + i,
            metadata={"ocr_engine": "claude_doc", "ocr_confidence": 100},
        ))

    return out if out else _extract_page_ocr(page, page_index, order_start)


def _extract_page_ocr(page, page_index: int, order_start: int) -> List[DocumentBlock]:
    """이미지 기반(스캔) PDF 페이지를 Tesseract OCR로 텍스트 추출.

    반환값: pdf_ocr_span 타입의 DocumentBlock 리스트.
    각 블록은 라인 단위로 묶이며, PDF 포인트 좌표의 bbox를 포함합니다.
    """
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "OCR 기능에는 pytesseract와 Pillow가 필요합니다.\n"
            "  pip install pytesseract Pillow\n"
            "  (시스템에 tesseract-ocr 바이너리도 필요합니다)"
        ) from exc

    import fitz  # type: ignore  # 이미 임포트되어 있지만 명시

    # ── 페이지를 고해상도 이미지로 렌더링 ───────────────────────────────
    mat = fitz.Matrix(_OCR_ZOOM, _OCR_ZOOM)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # ── Tesseract OCR (단어 + bounding box 포함 출력) ────────────────────
    try:
        ocr_data = pytesseract.image_to_data(
            img,
            output_type=pytesseract.Output.DICT,
            lang=_OCR_LANG,
            config="--psm 3",  # 자동 페이지 분할 (mixed layout 대응)
        )
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError(
            "tesseract 실행파일을 찾을 수 없습니다.\n"
            "Dockerfile에 tesseract-ocr 패키지를 추가하세요:\n"
            "  apt-get install tesseract-ocr tesseract-ocr-jpn tesseract-ocr-eng"
        ) from exc

    # ── 단어를 (block, par, line) 기준으로 묶어 라인 단위 블록 생성 ──────
    # Tesseract Output.DICT 구조: 동일 key는 같은 라인
    line_map: dict = {}
    for i in range(len(ocr_data["text"])):
        word = ocr_data["text"][i].strip()
        conf = int(ocr_data["conf"][i])
        # conf == -1 은 레이아웃 구분자, _OCR_MIN_CONF 미만은 노이즈
        if conf < _OCR_MIN_CONF or not word:
            continue

        key = (
            ocr_data["block_num"][i],
            ocr_data["par_num"][i],
            ocr_data["line_num"][i],
        )
        px_x  = ocr_data["left"][i]
        px_y  = ocr_data["top"][i]
        px_x2 = px_x + ocr_data["width"][i]
        px_y2 = px_y + ocr_data["height"][i]

        if key not in line_map:
            line_map[key] = {
                "words":  [],
                "x":  px_x,   "y":  px_y,
                "x2": px_x2,  "y2": px_y2,
                "confs": [],
            }
        else:
            line_map[key]["x2"] = max(line_map[key]["x2"], px_x2)
            line_map[key]["y2"] = max(line_map[key]["y2"], px_y2)

        line_map[key]["words"].append(word)
        line_map[key]["confs"].append(conf)

    # ── DocumentBlock 생성 ───────────────────────────────────────────────
    result: List[DocumentBlock] = []
    for idx, (_, ln) in enumerate(sorted(line_map.items())):
        text = " ".join(ln["words"]).strip()
        if not text:
            continue

        # 픽셀 좌표 → PDF 포인트 좌표 (zoom 역변환)
        x0 = ln["x"]  / _OCR_ZOOM
        y0 = ln["y"]  / _OCR_ZOOM
        x1 = ln["x2"] / _OCR_ZOOM
        y1 = ln["y2"] / _OCR_ZOOM

        line_height = y1 - y0
        font_size   = max(6.0, round(line_height * 0.75, 1))
        avg_conf    = sum(ln["confs"]) / len(ln["confs"])

        result.append(
            DocumentBlock(
                id=f"page-{page_index + 1}-ocr-{order_start + idx + 1}",
                type="pdf_ocr_span",
                sourceText=text,
                translatedText=None,
                pageNumber=page_index + 1,
                bbox=(x0, y0, x1, y1),
                fontSize=font_size,
                order=order_start + idx,
                metadata={
                    "ocr_confidence": round(avg_conf, 1),
                    "font": None,
                    "color": 0,
                },
            )
        )

    return result


def _extract_page_claude_vision(page, page_index: int, order_start: int) -> List[DocumentBlock]:
    """Claude Vision OCR + Tesseract bbox 하이브리드.

    1. Tesseract로 단락 bbox 탐지 (위치 정보)
    2. Claude Vision HTTP API로 전체 텍스트 추출 (텍스트 품질)
    3. bbox ↔ 텍스트 매핑 → pdf_ocr_span 블록 반환 (원위치 교체 지원)

    Tesseract 미설치 시: pdf_vision_span으로 폴백 (bbox 없음, 오버레이 출력)
    Claude Vision API 실패 시: Tesseract 결과로 폴백
    """
    import json
    import urllib.request

    import fitz  # type: ignore

    from app.provider_settings import provider_model as _pm, provider_secret

    api_key = provider_secret("anthropic", "api_key", "ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("Anthropic API 키 미설정 (config/providers.local.json)")
    vision_model = _pm("anthropic", "model", "ANTHROPIC_TRANSLATION_MODEL", _VISION_MODEL)

    # ── Step 1: Tesseract로 단락 bbox 탐지 ──────────────────────────────
    para_bboxes = _get_paragraph_bboxes(page)

    # ── Step 2: Claude Vision HTTP API 호출 ─────────────────────────────
    mat = fitz.Matrix(2.0, 2.0)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_b64 = base64.standard_b64encode(pix.tobytes("png")).decode()

    payload = json.dumps({
        "model": vision_model,
        "max_tokens": 4096,
        "messages": [{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {"type": "base64", "media_type": "image/png", "data": img_b64},
                },
                {
                    "type": "text",
                    "text": (
                        "Extract all text from this image exactly as it appears. "
                        "Rules:\n"
                        "- Output ONLY the raw text, no explanations or comments\n"
                        "- Preserve original language (Japanese, Chinese, Korean, etc.)\n"
                        "- Preserve numbers, punctuation, and symbols exactly\n"
                        "- For vertical Japanese text, output each column left-to-right\n"
                        "- Separate distinct text blocks with a blank line\n"
                        "- Do NOT translate, summarize, or add markdown formatting"
                    ),
                },
            ],
        }],
    }).encode("utf-8")

    try:
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            body = json.loads(resp.read().decode("utf-8"))
        ocr_text = body["content"][0]["text"].strip()
        logger.info("[Vision OCR] 페이지 %d 추출 완료: %d chars", page_index + 1, len(ocr_text))
    except Exception as exc:
        logger.error("[Vision OCR] 페이지 %d API 실패: %s", page_index + 1, exc)
        # API 실패 → Tesseract 결과로 폴백
        if para_bboxes:
            return _extract_page_ocr(page, page_index, order_start)
        raise RuntimeError(f"Claude Vision API 호출 실패: {exc}") from exc

    # ── Step 3: Claude Vision 텍스트를 단락 단위로 분할 ─────────────────
    cv_paras = [p.strip() for p in re.split(r"\n{2,}", ocr_text) if p.strip()]
    if not cv_paras:
        cv_paras = [ln.strip() for ln in ocr_text.splitlines() if ln.strip()]
    if not cv_paras:
        cv_paras = [ocr_text] if ocr_text else []

    # ── Tesseract bbox 없으면 pdf_vision_span으로 폴백 ───────────────────
    if not para_bboxes:
        result: List[DocumentBlock] = []
        for idx, para in enumerate(cv_paras):
            result.append(DocumentBlock(
                id=f"page-{page_index + 1}-vision-{order_start + idx + 1}",
                type="pdf_vision_span",
                sourceText=para,
                translatedText=None,
                pageNumber=page_index + 1,
                bbox=None,
                order=order_start + idx,
                metadata={"ocr_engine": "claude_vision", "ocr_confidence": 100},
            ))
        return result

    # ── Step 4: Tesseract bbox ↔ Claude Vision 텍스트 비례 매핑 ─────────
    # Tesseract 단락 수(n)에 맞게 Claude Vision 단락(m)을 균등 분배
    n, m = len(para_bboxes), len(cv_paras)
    out: List[DocumentBlock] = []
    for i, (x0, y0, x1, y1, _) in enumerate(para_bboxes):
        start = round(i * m / n)
        end   = round((i + 1) * m / n)
        text  = " ".join(cv_paras[start:end]).strip()
        if not text:
            continue
        font_size = max(6.0, round((y1 - y0) * 0.75, 1))
        out.append(DocumentBlock(
            id=f"page-{page_index + 1}-vision-{order_start + i + 1}",
            type="pdf_ocr_span",        # bbox 있음 → 원위치 교체
            sourceText=text,
            translatedText=None,
            pageNumber=page_index + 1,
            bbox=(x0, y0, x1, y1),
            fontSize=font_size,
            order=order_start + i,
            metadata={"ocr_engine": "claude_vision", "ocr_confidence": 100},
        ))

    return out if out else _extract_page_ocr(page, page_index, order_start)


def _get_paragraph_bboxes(page) -> List[tuple]:
    """Tesseract로 단락 레벨 bbox만 추출 (텍스트 위치 정보 전용).

    반환: [(x0, y0, x1, y1, tess_text), ...] — PDF 포인트 좌표
    pytesseract/Pillow 미설치 시 빈 리스트 반환.
    """
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        return []

    import fitz  # type: ignore

    mat = fitz.Matrix(_OCR_ZOOM, _OCR_ZOOM)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    try:
        ocr_data = pytesseract.image_to_data(
            img,
            output_type=pytesseract.Output.DICT,
            lang=_OCR_LANG,
            config="--psm 3",
        )
    except Exception as exc:
        logger.warning("[Vision OCR] Tesseract bbox 추출 실패: %s", exc)
        return []

    para_map: dict = {}
    for i in range(len(ocr_data["text"])):
        word = ocr_data["text"][i].strip()
        conf = int(ocr_data["conf"][i])
        if conf < _OCR_MIN_CONF or not word:
            continue

        key = (ocr_data["block_num"][i], ocr_data["par_num"][i])
        px, py = ocr_data["left"][i], ocr_data["top"][i]
        px2 = px + ocr_data["width"][i]
        py2 = py + ocr_data["height"][i]

        if key not in para_map:
            para_map[key] = {"x": px, "y": py, "x2": px2, "y2": py2, "words": [word]}
        else:
            p = para_map[key]
            p["x"], p["y"]   = min(p["x"], px),   min(p["y"], py)
            p["x2"], p["y2"] = max(p["x2"], px2), max(p["y2"], py2)
            p["words"].append(word)

    result = []
    for _, p in sorted(para_map.items()):
        x0, y0 = p["x"] / _OCR_ZOOM, p["y"] / _OCR_ZOOM
        x1, y1 = p["x2"] / _OCR_ZOOM, p["y2"] / _OCR_ZOOM
        result.append((x0, y0, x1, y1, " ".join(p["words"])))
    return result


def extract_docx_blocks(path: Path) -> List[DocumentBlock]:
    """python-docx로 단락 및 표 셀 텍스트를 추출.
    메타데이터에 단락/테이블 인덱스를 저장해 내보내기 시 원위치 교체에 사용."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "DOCX 추출에는 python-docx가 필요합니다: pip install python-docx"
        ) from exc

    doc = Document(path)
    blocks: List[DocumentBlock] = []
    order = 0

    # ── 일반 단락 ──────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        blocks.append(
            DocumentBlock(
                id=f"para-{para_idx}",
                type="docx_paragraph",
                sourceText=text,
                order=order,
                metadata={
                    "para_index": para_idx,
                    "style": para.style.name if para.style else "",
                    "alignment": str(para.alignment),
                },
            )
        )
        order += 1

    # ── 표 셀 ──────────────────────────────────────────────
    for tbl_idx, table in enumerate(doc.tables):
        seen_cells: set = set()          # 병합 셀 중복 방지
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_key = id(cell._tc)  # 병합 셀은 같은 _tc를 공유
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                text = cell.text.strip()
                if not text:
                    continue
                blocks.append(
                    DocumentBlock(
                        id=f"tbl-{tbl_idx}-r{row_idx}-c{col_idx}",
                        type="docx_table_cell",
                        sourceText=text,
                        order=order,
                        metadata={
                            "table_index": tbl_idx,
                            "row_index": row_idx,
                            "col_index": col_idx,
                        },
                    )
                )
                order += 1

    return blocks


def extract_doc_blocks(path: Path) -> List[DocumentBlock]:
    """구형 .doc 바이너리 형식 — antiword CLI로 텍스트를 추출.
    서식 정보는 손실되며, 내보내기 시 새 DOCX 파일로 생성됩니다."""
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=60,
        )
    except FileNotFoundError:
        raise RuntimeError(
            "antiword가 설치되어 있지 않습니다. "
            "Docker 환경에서는 Dockerfile에 'antiword' 패키지를 추가하세요."
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("antiword 실행 시간 초과")

    if result.returncode != 0:
        raise RuntimeError(
            f"antiword 실행 실패 (code={result.returncode}): {result.stderr.strip() or '알 수 없는 오류'}"
        )

    text = result.stdout
    parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not parts and text.strip():
        parts = [text.strip()]

    blocks = []
    for idx, part in enumerate(parts):
        blocks.append(
            DocumentBlock(
                id=f"para-{idx}",
                type="docx_paragraph",   # export 시 docx 단락으로 취급
                sourceText=part,
                order=idx,
                metadata={"fileType": "doc", "para_index": idx},
            )
        )
    return blocks


def extract_xlsx_blocks(path: Path) -> List[DocumentBlock]:
    """openpyxl로 시트별 셀 텍스트를 추출.
    메타데이터에 시트명·셀 좌표를 저장해 내보내기 시 원위치 교체에 사용."""
    try:
        import openpyxl  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "XLSX 추출에는 openpyxl이 필요합니다: pip install openpyxl"
        ) from exc

    wb = openpyxl.load_workbook(path, data_only=True)
    blocks: List[DocumentBlock] = []
    order = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                raw = cell.value
                if raw is None:
                    continue
                text = str(raw).strip()
                if not text:
                    continue
                blocks.append(
                    DocumentBlock(
                        id=f"sheet-{sheet_name}-{cell.coordinate}",
                        type="xlsx_cell",
                        sourceText=text,
                        sheetName=sheet_name,
                        order=order,
                        metadata={
                            "sheet": sheet_name,
                            "coordinate": cell.coordinate,
                            "row": cell.row,
                            "col": cell.column,
                            "number_format": cell.number_format or "General",
                        },
                    )
                )
                order += 1

    return blocks


def extract_xls_blocks(path: Path) -> List[DocumentBlock]:
    """구형 .xls 바이너리 형식 — xlrd로 텍스트를 추출.
    내보내기 시 새 XLSX 파일로 생성됩니다."""
    try:
        import xlrd  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "XLS 추출에는 xlrd가 필요합니다: pip install 'xlrd>=1.2,<2'"
        ) from exc

    wb = xlrd.open_workbook(str(path))
    blocks: List[DocumentBlock] = []
    order = 0

    for sheet_idx in range(wb.nsheets):
        ws = wb.sheet_by_index(sheet_idx)
        sheet_name = ws.name
        for row_idx in range(ws.nrows):
            for col_idx in range(ws.ncols):
                cell = ws.cell(row_idx, col_idx)
                if cell.ctype == xlrd.XL_CELL_EMPTY:
                    continue
                text = str(cell.value).strip()
                if not text:
                    continue
                # openpyxl 호환 좌표 문자열 생성 (A1, B2, ...)
                col_letter = _col_letter(col_idx)
                coordinate = f"{col_letter}{row_idx + 1}"
                blocks.append(
                    DocumentBlock(
                        id=f"sheet-{sheet_name}-{coordinate}",
                        type="xlsx_cell",   # export 시 xlsx_cell로 취급
                        sourceText=text,
                        sheetName=sheet_name,
                        order=order,
                        metadata={
                            "sheet": sheet_name,
                            "coordinate": coordinate,
                            "row": row_idx + 1,
                            "col": col_idx + 1,
                        },
                    )
                )
                order += 1

    return blocks


def _col_letter(col_idx: int) -> str:
    """0-based column index → Excel 열 문자 (0→A, 25→Z, 26→AA ...)"""
    result = ""
    n = col_idx
    while True:
        result = chr(ord("A") + n % 26) + result
        n = n // 26 - 1
        if n < 0:
            break
    return result
